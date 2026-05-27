import docx

def process_docx(input_path, output_path, translator):
    """
    Traite un document docx, extrait le texte, le traduit en une passe, 
    et reconstruit le document.
    """
    print(f"Chargement du document DOCX: {input_path}")
    doc = docx.Document(input_path)
    
    texts_to_translate = []
    
    # 1. Extraction (paragraphes et tableaux)
    for para in doc.paragraphs:
        if para.text.strip():
            texts_to_translate.append(para.text)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        texts_to_translate.append(para.text)
                        
    if not texts_to_translate:
        print("Aucun texte à traduire trouvé dans le DOCX.")
        doc.save(output_path)
        return
        
    print(f"{len(texts_to_translate)} blocs de texte extraits.")
    
    # 2. Traduction de l'ensemble (le translator gère le batching et multiprocessing)
    translated_texts = translator.translate_texts(texts_to_translate)
    
    # 3. Réinsertion (l'ordre d'itération est déterministe)
    text_index = 0
    
    for para in doc.paragraphs:
        if para.text.strip():
            _replace_paragraph_text(para, translated_texts[text_index])
            text_index += 1
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        _replace_paragraph_text(para, translated_texts[text_index])
                        text_index += 1
                        
    print(f"Sauvegarde du DOCX traduit: {output_path}")
    doc.save(output_path)

def _replace_paragraph_text(paragraph, new_text):
    """
    Remplace le texte d'un paragraphe docx de manière robuste.
    Conserve le premier 'run' pour garder la police et taille de base,
    et vide les autres.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
        
    first_run = paragraph.runs[0]
    first_run.text = new_text
    
    for i in range(1, len(paragraph.runs)):
        paragraph.runs[i].text = ""
